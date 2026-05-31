import os
import uuid
import logging
from pathlib import Path
from typing import List
from fastapi import UploadFile, HTTPException
from sqlalchemy.orm import Session

from ..models.document_model import Document
from ..models.business_model import Business
from ..schemas.document_schema import DocumentResponse
from ..config import UPLOAD_DIR

import argostranslate.translate
from langdetect import detect

from pdfminer.high_level import extract_text as pdf_extract
from docx import Document as DocxDocument

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB

logger = logging.getLogger("document_service")
logger.setLevel(logging.INFO)

# ----------------------------
# TRANSLATION
# ----------------------------
def translate_text(text: str, source_lang: str, target_lang: str) -> str:
    if not text or source_lang == target_lang:
        return text

    try:
        return argostranslate.translate.translate(text, source_lang, target_lang)
    except Exception as e:
        logger.error(f"Argos failed: {e}")
        return text  # fallback (free + safe)


# ----------------------------
# TEXT EXTRACTION
# ----------------------------
def extract_text(file_path: str, ext: str) -> str:
    ext = ext.lower()

    try:
        # PDF
        if ext == ".pdf":
            return pdf_extract(file_path)

        # DOCX
        elif ext == ".docx":
            doc = DocxDocument(file_path)
            return "\n".join([p.text for p in doc.paragraphs])

        # TXT
        elif ext in [".txt", ".md"]:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()

        # Images → (OCR can be added later if needed)
        else:
            return ""

    except Exception as e:
        logger.error(f"Text extraction failed: {e}")
        return ""


# ----------------------------
# SAVE FILE
# ----------------------------
async def save_file(
        db: Session,
        file: UploadFile,
        business: Business,
        lang: str
) -> DocumentResponse:

    os.makedirs(UPLOAD_DIR, exist_ok=True)

    if not file.filename:
        raise HTTPException(status_code=400, detail="Invalid file")

    ext = Path(file.filename).suffix
    filename = f"{uuid.uuid4()}{ext}"
    path = os.path.join(UPLOAD_DIR, filename)

    size = 0

    try:
        with open(path, "wb") as f:
            while chunk := await file.read(1024 * 1024):
                size += len(chunk)

                if size > MAX_FILE_SIZE:
                    os.remove(path)
                    raise HTTPException(status_code=413, detail="File too large")

                f.write(chunk)

    except Exception:
        raise HTTPException(status_code=500, detail="Failed to save file")

    # ✅ Extract text once
    extracted_text = extract_text(path, ext)

    # ✅ Detect language automatically
    detected_lang = None
    if extracted_text:
        try:
            detected_lang = detect(extracted_text)
        except:
            detected_lang = lang

    try:
        document = Document(
            business_id=business.id,
            file_path=path,
            file_type=file.content_type,
            file_size=size,
            language=detected_lang or lang,
            extracted_text=extracted_text  # ✅ store text
        )

        db.add(document)
        db.commit()
        db.refresh(document)

    except Exception as e:
        db.rollback()
        if os.path.exists(path):
            os.remove(path)
        raise HTTPException(status_code=500, detail="DB error")

    return DocumentResponse.from_orm(document)


# ----------------------------
# GET DOCUMENTS (SMART CACHE)
# ----------------------------
def get_documents(
        db: Session,
        business: Business,
        target_lang: str
) -> List[DocumentResponse]:

    docs = db.query(Document).filter(
        Document.business_id == business.id
    ).all()

    results = []

    for doc in docs:

        # ✅ If already in target language → return directly
        if doc.language == target_lang:
            results.append(doc)
            continue

        # ✅ Check if translation already exists in DB
        existing = db.query(Document).filter(
            Document.business_id == business.id,
            Document.translated_from == doc.language,
            Document.language == target_lang
        ).first()

        if existing:
            results.append(existing)
            continue

        # ✅ Translate only extracted text (NOT file)
        translated_text = translate_text(
            doc.extracted_text,
            doc.language,
            target_lang
        )

        # ✅ Save translated version
        new_doc = Document(
            business_id=business.id,
            file_path=doc.file_path,
            file_type=doc.file_type,
            file_size=doc.file_size,
            language=target_lang,
            translated_from=doc.language,
            extracted_text=translated_text
        )

        db.add(new_doc)
        results.append(new_doc)

    db.commit()

    return [DocumentResponse.from_orm(d) for d in results]


# ----------------------------
# DELETE DOCUMENT
# ----------------------------
def delete_document(
        db: Session,
        document_id: int,
        business: Business
):

    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.business_id == business.id
    ).first()

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if os.path.exists(doc.file_path):
        os.remove(doc.file_path)

    db.delete(doc)
    db.commit()